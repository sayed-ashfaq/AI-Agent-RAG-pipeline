import os
import tempfile
from pathlib import Path
from datetime import datetime
from typing import List, Tuple, Dict, Any
import pandas as pd
import pymupdf  # PyMuPDF
import base64
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from src.utils.model_loader import ModelLoader
from src.utils.config_loader import load_config
from langchain_core.messages import HumanMessage
from langchain_community.document_loaders import (
    UnstructuredMarkdownLoader,
    UnstructuredPowerPointLoader,
    UnstructuredFileLoader
)
import sqlite3
from src.tools.rag_agent.schemas import FileType, DocumentInfo, BatchUploadResponse
from custom_logger import GLOBAL_LOGGER as logger
from exception_handler.agent_exceptions import FileProcessingError

class DocumentService:
    def __init__(self):
        self.config = load_config()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size= self.config["retriever"]["chunk_size"],
            chunk_overlap= self.config["retriever"]['chunk_overlap'],
        )
        self.vision_llm = ModelLoader().load_llm()

    def process_multiple_files(self, files: List[Tuple[str, bytes, str]]) -> Tuple[List[Document], List[DocumentInfo]]:
        """Process multiple files at once"""
        logger.info("Initiated process_multiple_files")
        all_documents = []
        document_infos = []

        for filename, file_content, file_type in files:
            try:
                # Save temp file
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_type}") as tmp_file:
                    tmp_file.write(file_content)
                    tmp_file_path = tmp_file.name

                # Process based on type
                documents = self.process_single_file(tmp_file_path, FileType(file_type))

                # Create document info
                doc_info = DocumentInfo(
                    document_id=f"{Path(filename).stem}_{hash(filename)}",
                    filename=filename,
                    file_type=FileType(file_type),
                    processed_at=datetime.now(),
                    chunk_count=len(documents),
                    # has_images=metadata.get("has_images", False),
                    # image_count=metadata.get("extracted_images", 0)
                )

                all_documents.extend(documents)
                document_infos.append(doc_info)

                # Cleanup
                os.unlink(tmp_file_path)

            except Exception as e:
                print(f"Failed to process {filename}: {e}")
                continue

        return all_documents, document_infos

    def process_single_file(self, file_path: str, file_type: FileType) -> List[Document]:
        """Process a single file based on its type"""
        if file_type == FileType.PDF:
            return self._process_pdf(file_path)
        elif file_type == FileType.IMAGE:
            return self._process_image(file_path)
        elif file_type == FileType.CSV:
            return self._process_csv(file_path)
        elif file_type == FileType.XLSX:
            return self._process_excel(file_path)
        elif file_type == FileType.TXT:
            return self._process_text(file_path)
        elif file_type == FileType.MARKDOWN:
            return self._process_markdown(file_path)
        elif file_type == FileType.DOCX:
            return self._process_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _process_pdf_multimodal(self, file_path: str) -> List[Document]:
        """Process PDF with both text and images"""
        documents = []
        metadata = {"pages": 0, "has_images": False, "extracted_images": 0}

        # Use PyMuPDF for comprehensive extraction
        pdf_document = pymupdf.open(file_path)
        metadata["pages"] = len(pdf_document)

        full_text = ""
        image_descriptions = []

        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)

            # Extract text
            page_text = page.get_text()
            if page_text.strip():
                full_text += f"\n[Page {page_num + 1}]\n{page_text}"

            # Extract images
            image_list = page.get_images()
            if image_list:
                metadata["has_images"] = True

                for img_index, img in enumerate(image_list):
                    try:
                        # Extract image
                        xref = img[0]
                        pix = pymupdf.Pixmap(pdf_document, xref)

                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            img_data = pix.tobytes("png")
                            img_base64 = base64.b64encode(img_data).decode()

                            # Describe image using vision model
                            description = self._describe_image(img_base64, page_num + 1, img_index + 1)

                            if description:
                                image_descriptions.append({
                                    "page": page_num + 1,
                                    "index": img_index + 1,
                                    "description": description
                                })
                                metadata["extracted_images"] += 1

                                # Add to text content
                                full_text += f"\n[Image {img_index + 1} on Page {page_num + 1}]: {description}\n"

                        pix = None
                    except Exception as e:
                        continue

        pdf_document.close()

        # Create text documents
        if full_text.strip():
            text_docs = self.text_splitter.create_documents(
                [full_text],
                metadatas=[{
                    "source": file_path,
                    "type": "pdf",
                    "pages": metadata["pages"],
                    "has_images": metadata["has_images"]
                }]
            )
            documents.extend(text_docs)

        # Create separate documents for significant images
        for img_desc in image_descriptions:
            if len(img_desc["description"]) > 100:
                img_doc = Document(
                    page_content=f"Image from Page {img_desc['page']}: {img_desc['description']}",
                    metadata={
                        "source": file_path,
                        "type": "pdf_image",
                        "page": img_desc["page"]
                    }
                )
                documents.append(img_doc)

        return documents

    def _process_image(self, file_path: str) -> List[Document]:
        """Process standalone images"""
        # Convert image to base64
        with open(file_path, "rb") as image_file:
            img_data = image_file.read()
            img_base64 = base64.b64encode(img_data).decode()

        # Get image description
        description = self._describe_image(img_base64, filename=Path(file_path).name)

        document = Document(
            page_content=f"Image: {Path(file_path).name}\nDescription: {description}",
            metadata={
                "source": file_path,
                "type": "image",
                "filename": Path(file_path).name
            }
        )

        # metadata = {"has_images": True, "extracted_images": 1}
        return [document]


    def _process_pdf(self, file_path: str) -> List[Document]:
        """Process PDF files (with multimodal support)"""
        try:
            documents= []
            metadata= {"pages": 0, "has_images": False}

            pdf_document = pymupdf.open(file_path)
            metadata["pages"] = len(pdf_document)

            # extract text
            full_text = ""
            for page_num in range(len(pdf_document)):
                page= pdf_document.load_page(page_num)

                page_text= page.get_text()
                if page_text.strip():
                    full_text += f"\n[Page {page_num + 1}]\n{page_text}"
            pdf_document.close()

            if full_text.strip():
                text_docs= self.text_splitter.create_documents(
                    [full_text],
                    metadatas=[{
                        "source": file_path,
                        'type': 'pdf',
                        'pages': metadata["pages"],
                        'has_images': metadata["has_images"]
                    }]
                )
                documents.extend(text_docs)

                logger.info(f"Processed PDF: {len(documents)} chunks from {metadata['pages']} pages")
                return documents
            return []

        except Exception as e:
            raise FileProcessingError(f"PDF processing failed: {e}")


    def _process_csv(self, file_path: str) -> List[Document]:
        """Process CSV files with data analysis"""
        df = pd.read_csv(file_path)

        # Create comprehensive content
        content = f"CSV File Analysis\n"
        content += f"Filename: {Path(file_path).name}\n"
        content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n\n"

        # Column information
        content += "Columns:\n"
        for col in df.columns:
            content += f"- {col}: {df[col].dtype}\n"

        # Sample data (first 20 rows)
        content += f"\nSample Data (first 20 rows):\n"
        content += df.head(20).to_string(index=False)

        # Statistical summary for numeric columns
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            content += f"\n\nStatistical Summary:\n"
            content += df[numeric_cols].describe().to_string()

        # Missing values
        missing_values = df.isnull().sum()
        if missing_values.sum() > 0:
            content += f"\n\nMissing Values:\n"
            for col, missing in missing_values.items():
                if missing > 0:
                    content += f"- {col}: {missing} missing values\n"

        documents = self.text_splitter.create_documents(
            [content],
            metadatas=[{"source": file_path, "type": "csv", "rows": len(df), "columns": len(df.columns)}]
        )

        metadata = {"rows": len(df), "columns": len(df.columns)}
        return documents

    def _process_excel(self, file_path: str) -> List[Document]:
        """Process Excel files with multiple sheets"""
        documents = []
        xl_file = pd.ExcelFile(file_path)
        metadata = {"sheets": len(xl_file.sheet_names), "total_rows": 0}

        for sheet_name in xl_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            # Create sheet content
            sheet_content = f"Excel Sheet: {sheet_name}\n"
            sheet_content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n\n"

            # Column info
            sheet_content += "Columns:\n"
            for col in df.columns:
                sheet_content += f"- {col}: {df[col].dtype}\n"

            # Sample data
            sheet_content += f"\nSample Data:\n"
            sheet_content += df.head(15).to_string(index=False)

            # Statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                sheet_content += f"\n\nStatistical Summary:\n"
                sheet_content += df[numeric_cols].describe().to_string()

            sheet_docs = self.text_splitter.create_documents(
                [sheet_content],
                metadatas=[{"source": file_path, "type": "excel", "sheet": sheet_name}]
            )
            documents.extend(sheet_docs)
            metadata["total_rows"] += len(df)

        return documents

    def _process_text(self, file_path: str) -> List[Document]:
        """Process plain text files"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()

        documents = self.text_splitter.create_documents(
            [content],
            metadatas=[{"source": file_path, "type": "txt"}]
        )

        metadata = {"characters": len(content)}
        return documents

    def _process_markdown(self, file_path: str) -> List[Document]:
        """Process Markdown files using LangChain's UnstructuredMarkdownLoader"""
        try:
            # Load markdown with unstructured
            loader = UnstructuredMarkdownLoader(file_path, mode="elements")
            elements = loader.load()

            metadata = {
                "total_elements": len(elements),
                "has_code_blocks": False,
                "has_tables": False,
                "has_lists": False
            }

            # Combine elements into structured content
            content_parts = []
            for element in elements:
                element_metadata = element.metadata
                category = element_metadata.get("category", "Text")

                # Track content types
                if category == "CodeSnippet":
                    metadata["has_code_blocks"] = True
                elif category == "Table":
                    metadata["has_tables"] = True
                elif category == "ListItem":
                    metadata["has_lists"] = True

                content_parts.append(f"[{category}] {element.page_content}")

            full_content = "\n\n".join(content_parts)

            # Split into chunks
            documents = self.text_splitter.create_documents(
                [full_content],
                metadatas=[{
                    "source": file_path,
                    "type": "markdown",
                    "total_elements": metadata["total_elements"],
                    "has_code_blocks": metadata["has_code_blocks"],
                    "has_tables": metadata["has_tables"]
                }]
            )

            logger.info(f"Processed Markdown: {len(documents)} chunks, {metadata['total_elements']} elements")
            return documents

        except Exception as e:
            raise FileProcessingError(f"Markdown processing failed: {e}")


    def _process_docx(self, file_path: str) -> List[Document]:
        """Process Word documents"""
        doc = DocxDocument(file_path)

        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = '\n\n'.join(paragraphs)

        # Extract tables
        table_content = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(' | '.join(row_data))
            if table_data:
                table_content.append('\n'.join(table_data))

        if table_content:
            content += '\n\n[TABLES]\n' + '\n\n'.join(table_content)

        documents = self.text_splitter.create_documents(
            [content],
            metadatas=[{"source": file_path, "type": "docx"}]
        )

        metadata = {"paragraphs": len(paragraphs), "tables": len(doc.tables)}
        return documents

    def _describe_image(self, img_base64: str, page_num: int = None, img_index: int = None, filename: str = None) -> str:
        """Generate image description using vision model"""
        try:
            location = f"page {page_num}, image {img_index}" if page_num else f"file {filename}"

            message_content = [
                {
                    "type": "text",
                    "text": f"""Analyze this image from a business document ({location}).
                    Focus on:
                    - Charts, graphs, tables, diagrams
                    - Key data points, trends, metrics
                    - Business-relevant visual information
                    - Any text or labels visible

                    Provide a clear, detailed description of what's shown."""
                },
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/png;base64,{img_base64}"}
                }
            ]

            response = self.vision_llm.invoke([HumanMessage(content=message_content)])
            return response.content.strip()

        except Exception as e:
            return f"Image from {location} (description unavailable: {str(e)})"

    def process_database_query(self, db_connection: str, query: str, table_name: str = None) -> Tuple[List[Document], Dict[str, Any]]:
        """Process database queries and convert to documents"""
        documents = []
        metadata = {}

        try:
            # Connect to database (supporting SQLite for now)
            conn = sqlite3.connect(db_connection)

            if query:
                # Execute custom query
                df = pd.read_sql_query(query, conn)
                content = f"Database Query Results\n"
                content += f"Query: {query}\n"
                content += f"Results: {len(df)} rows, {len(df.columns)} columns\n\n"
            else:
                # Get table info if no query provided
                if table_name:
                    df = pd.read_sql_query(f"SELECT * FROM {table_name} LIMIT 1000", conn)
                    content = f"Database Table: {table_name}\n"
                    content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n\n"
                else:
                    raise ValueError("Either query or table_name must be provided")

            # Add data content
            content += "Column Information:\n"
            for col in df.columns:
                content += f"- {col}: {df[col].dtype}\n"

            content += f"\nSample Data:\n"
            content += df.head(50).to_string(index=False)

            # Statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                content += f"\n\nStatistical Summary:\n"
                content += df[numeric_cols].describe().to_string()

            documents = self.text_splitter.create_documents(
                [content],
                metadatas=[{"source": db_connection, "type": "database", "table": table_name or "query"}]
            )

            metadata = {"rows": len(df), "columns": len(df.columns)}
            conn.close()

        except Exception as e:
            raise Exception(f"Database processing failed: {e}")

        return documents, metadata